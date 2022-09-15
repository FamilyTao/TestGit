
from ast import pattern
from contextlib import nullcontext
from enum import Flag, unique
from fileinput import filename
from multiprocessing.reduction import duplicate
from os import dup
from re import T
from time import sleep
from unittest.mock import patch
from xml.dom.minidom import Document
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font,Alignment, colors
from openpyxl.styles import PatternFill
import copy
import openpyxl
import shutil
import os

COLOR_INDEX = (
    '000000','FF0000','00FF00'
)

BLACK = COLOR_INDEX[0]
RED = COLOR_INDEX[1]
GREEN = COLOR_INDEX[2]

class SplitMultiSheet:
    """复制带格式"""

    def split_sheet(self, wb: openpyxl.Workbook, sn: str, newFileName:str):
        """
        :param wb:  excel表格
        :param sn:  sheet_name
        :return:
        """
        # _sn: sheet_name, ws: Worksheet
        for _sn, ws in zip(wb.sheetnames, wb):
            if _sn != sn:
                wb.remove(ws)
        # 保存文件、关闭文件
        wb.save(newFileName)
        shutil.move(newFileName,'output')
        wb.close()

    def main(self, file_path: str, newFileName: str):
        wb = openpyxl.load_workbook(filename=file_path)
        for sn in wb.sheetnames:
            # 深拷贝
            new_wb = copy.deepcopy(wb)
            self.split_sheet(wb=new_wb, sn=sn, newFileName = newFileName)
        return new_wb


class OfficeExcel:
    def __init__(self, pathFileName:str):
        self.pathFilename = pathFileName
        self.path = os.path.dirname(pathFileName)
        self.fileName = os.path.basename(pathFileName)
        self.currentSheetName = nullcontext
        self.wb = nullcontext
        self.sheet = nullcontext

    def create_excel(self):
        from openpyxl import Workbook
        wb = Workbook()
        wb.create_sheet()
        wb.save(self.fileName)
        return load_workbook(self.fileName)

    # 打开excel
    def open(self):
        try:
            self.wb = load_workbook(self.pathFilename)
        except FileNotFoundError:
            self.wb = self.create_excel()
            print(f'[warn] 文件不存在,创建文件{self.fileName}')
        else:
            print('[log] open {0} success'.format(self.fileName))
            return self.wb
    
    # 获取所有工作簿名单
    def show_all_sheetnames(self):
        return self.wb.sheetnames

    def select_sheet(self,sheetName):
        self.currentSheetName = sheetName
        self.sheet = self.wb[sheetName]
        print('[log] 当前{0}工作簿被选中'.format(sheetName))
        return self.sheet 
    
    def get_value_by_pos(self,row,col):
        return self.sheet.cell(row,col).value

    def get_value_by_row(self,row):
        values = []
        for cell in self.sheet[row]:
            if(None != cell.value):
                values.append(cell.value)
        return values

    def get_value_by_col(self,col:int):
        values = []
        for cell in self.sheet[get_column_letter(col)]:
            if(None != cell.value):
                values.append(cell.value)
        return values
    
    def get_value_by_row_scope(self,row_start,row_end):
        values = {}
        for row in self.sheet[row_start:row_end]:
            for cell in row:
                if(None != cell.value):
                    values.setdefault(str(cell.row),[]).append(cell.value)
        return values

    def get_value_by_col_scope(self,col_start,col_end):
        values = {}
        for col in self.sheet[get_column_letter(col_start):get_column_letter(col_end)]:
            for cell in col:
                if(None != cell.value):
                    values.setdefault(str(cell.column),[]).append(cell.value)
        return values

    # 行列切片
    def get_value_by_chip(self,row_start,row_end,col_start,col_end):
        values = []
        for row in self.sheet.iter_rows(min_row=row_start,max_row=row_end,min_col=col_start, max_col=col_end):
            for cell in row:
                if(None != cell.value):
                    values.append(cell.value)
        return values

    # 获取所有数据
    def get_value_by_sheet(self):
        values = []
        for row in self.sheet.rows:
            for cell in row:
                if(None != cell.value):
                    values.append(cell.value)
        return values
    
    # 根据坐标设置数据
    def set_value_by_pos(self,row,col,data):
        self.sheet.cell(row,col).value = data
        self.wb.save(self.fileName)
        shutil.move(self.fileName,'output')

    # 设置一行数据
    def set_value_by_row(self,row,data):
        for i,d in enumerate(data):
            self.sheet.cell(row,i+1).value = d
        self.wb.save(self.fileName)
        shutil.move(self.fileName,'output')

    # 设置一列数据
    def set_value_by_col(self,col,data):
        for i,d in enumerate(data):
            self.sheet.cell(i+1,col).value = d
        self.wb.save(self.fileName)

    # 设置整个表格数据(一行一行设置)
    def set_value_by_sheet(self,data):
        for row in data:
            self.sheet.append(row)
        self.wb.save(self.fileName)
    
    def set_cell_highlight(self,row,col,fontSize):
        bak_color = PatternFill('solid',fgColor='FFDEAD')
        fontConfig = Font(size=fontSize,color=RED)
        self.sheet.cell(row,col).font = fontConfig
        self.sheet.cell(row,col).fill = bak_color 
    
    def highlight_duplicate_by_col(self,col,colname:list):
        # 重复数据index
        duplicateIndex = []
        # 非重复数据
        noneDuolicateData = []
        bak_color = PatternFill('solid',fgColor='FFDEAD')

        # 先获取重复数据和非重复数据
        for i,cell in enumerate(self.sheet[get_column_letter(col)]):
            if cell.value not in noneDuolicateData:
                noneDuolicateData.append(cell.value)
            else:
                duplicateIndex.append(i)

        # 重复数据高亮
        for i,row in enumerate(self.sheet.rows):
            if i in duplicateIndex:
                for cell in row:
                    if(cell.column == col):
                        cell.fill = bak_color
        self.wb.save(self.fileName)
        
        # 去重数据到新的excel
        # excel拷贝
        duplicate = DuplicateData()
        duplicate.drop_duplicate_data(self.fileName,self.currentSheetName,colname)

import pandas as pd # pip3 install pandas
class DuplicateData:
    def drop_duplicate_data(self, filepath:str, shname:str, colnames:list):
        frame = pd.read_excel(filepath,sheet_name=shname)
        data = pd.DataFrame(frame)
        # print(data['姓名'].unique()) #取唯一值
        # print(data['姓名'].value_counts()) #重复次数
      
        # 删除重复
            # subset :列
            # keep : 'first' ,'last' ,False(不显示重复数据)
            # inplace 是否需要副本
        
        data.drop_duplicates(subset = colnames,keep='first',inplace=True)
        data.to_excel(filepath.replace('.xlsx','_去重数据.xlsx'))
        # 提取重复
    def duolicate_data(self,filepath,colnames:list):
        frame = pd.read_excel(filepath+'.xlsx')
        data = pd.DataFrame(frame)
        dupindex = data.duplicated(subset=colnames)
        data[dupindex].to_excel(filepath+'_重复数据.xlsx')
        return data[dupindex]
    '''案例'''
    # duplicate = DuplicateData()
    # duplicate.drop_duplicate_data('xlsx/5',['姓名'])
    # duplicate.duolicate_data('xlsx/5',['姓名'])

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
class OfficeWord:
   
    def __init__(self,pathFilename:str):
        self.pathFilename = pathFilename
        self.path = os.path.dirname(pathFilename)
        self.fileName = os.path.basename(pathFilename)
        self.doc = Document(self.pathFilename)

    def title_level(self,level:int):
        return 'Heading ' + str(level)
        
    def get_titles_by_level(self,level:title_level):
        titles = []
        for p in self.doc.paragraphs:
            if p.style.name == level:
                titles.append(p.text)
        return titles

    def get_all_titles(self):
        import re
        titles = []
        for p in self.doc.paragraphs:
            if re.match('^Heading \d+$',p.style.name):
                titles.append(p.text)
        return titles

    def get_all_maintext(self):
        text = ''
        for p in self.doc.paragraphs:
            if p.style.name == 'Normal' 'S-R Communciation':
                text += p.text
        return text

    def get_text_by_titlename(self,title:str,level:int):
        text = ''
        Flag = False
        for p in self.doc.paragraphs:
            
            if(Flag and (p.style.name == self.title_level(level=level) or p.style.name < self.title_level(level=level))):
                break

            if(p.style.name == self.title_level(level=level) and p.text == title):
                Flag = True
                continue

            if(Flag):
                text += p.text
        return text
    
    def get_regular_string_by_title(self,regular:str,title:str,level:int):
        ret = re.findall(regular,self.get_text_by_titlename(title,level))
        return ret

    # def get_string_by_title(self):
        
        
        # for p in doc.paragraphs:
        #     print(p.text)

    

import re
if __name__ == '__main__':
    path = 'Test/[PH-PRD-QC-002-2022]_ArchitecturalDesign_RTE'
    officeword = OfficeWord(path)
    officeword.get_titles_by_level(officeword.title_level(1))
    officeword.get_all_titles()
    
    # print(officeword.get_regular_string_by_title('SA_R2_\d*','Data Transformation',2)) 
